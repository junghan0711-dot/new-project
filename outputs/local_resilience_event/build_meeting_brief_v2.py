from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT = "/Users/junghanchiu/Documents/New project/outputs/local_resilience_event/地方韌性跨域共創發表會_世界咖啡館會議準備文本_v2.docx"
BLUE="2E74B5"; DARK="1F4D78"; INK="243746"; MUTED="6B7280"; LIGHT="E8EEF5"; PALE="F4F6F9"; GOLD="A46B12"

doc=Document(); sec=doc.sections[0]
sec.page_width,sec.page_height=Inches(8.5),Inches(11)
sec.top_margin=sec.bottom_margin=sec.left_margin=sec.right_margin=Inches(1)
sec.header_distance=sec.footer_distance=Inches(.492)

def setfont(run,size=11,bold=False,color="222222"):
    run.font.name="STHeiti"; run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"),"STHeiti")
    run.font.size=Pt(size); run.bold=bold; run.font.color.rgb=RGBColor.from_string(color)

normal=doc.styles["Normal"]; normal.font.name="STHeiti"; normal._element.rPr.rFonts.set(qn("w:eastAsia"),"STHeiti"); normal.font.size=Pt(11)
normal.paragraph_format.space_after=Pt(6); normal.paragraph_format.line_spacing=1.25
for name,size,color,before,after in [("Heading 1",16,BLUE,18,10),("Heading 2",13,BLUE,14,7),("Heading 3",12,DARK,10,5)]:
    st=doc.styles[name]; st.font.name="STHeiti"; st._element.rPr.rFonts.set(qn("w:eastAsia"),"STHeiti"); st.font.size=Pt(size); st.font.bold=True; st.font.color.rgb=RGBColor.from_string(color)
    st.paragraph_format.space_before=Pt(before); st.paragraph_format.space_after=Pt(after); st.paragraph_format.keep_with_next=True
for name in ["List Bullet","List Number"]:
    st=doc.styles[name]; st.font.name="STHeiti"; st._element.rPr.rFonts.set(qn("w:eastAsia"),"STHeiti"); st.font.size=Pt(11); st.paragraph_format.space_after=Pt(4); st.paragraph_format.line_spacing=1.2

def shade(cell,fill):
    pr=cell._tc.get_or_add_tcPr(); sh=OxmlElement("w:shd"); sh.set(qn("w:fill"),fill); pr.append(sh)

def cell(cell,text,bold=False,fill=None,color="222222",center=False,size=9.3):
    cell.text=""; p=cell.paragraphs[0]; p.paragraph_format.space_after=Pt(0); p.paragraph_format.line_spacing=1.12
    if center:p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    setfont(p.add_run(str(text)),size,bold,color); cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
    pr=cell._tc.get_or_add_tcPr(); mar=OxmlElement("w:tcMar")
    for side,val in [("top",90),("bottom",90),("start",120),("end",120)]:
        n=OxmlElement("w:"+side); n.set(qn("w:w"),str(val)); n.set(qn("w:type"),"dxa"); mar.append(n)
    pr.append(mar)
    if fill:shade(cell,fill)

def table(headers,rows,widths):
    t=doc.add_table(rows=1,cols=len(headers)); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=False
    for i,h in enumerate(headers): cell(t.rows[0].cells[i],h,True,LIGHT,DARK,True); t.rows[0].cells[i].width=Inches(widths[i])
    t.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    for row in rows:
        cs=t.add_row().cells
        for i,v in enumerate(row): cell(cs[i],v,False,None,"222222",i==0); cs[i].width=Inches(widths[i])
    doc.add_paragraph().paragraph_format.space_after=Pt(0); return t

def callout(label,text,fill=PALE):
    t=doc.add_table(rows=1,cols=1); t.autofit=False; t.columns[0].width=Inches(6.5); c=t.cell(0,0); shade(c,fill)
    c.text=""; p=c.paragraphs[0]; p.paragraph_format.space_after=Pt(0); p.paragraph_format.line_spacing=1.2; setfont(p.add_run(label+"｜"),11,True,DARK); setfont(p.add_run(text),11)
    doc.add_paragraph().paragraph_format.space_after=Pt(0)

def bullet(text):
    p=doc.add_paragraph(style="List Bullet"); setfont(p.add_run(text)); return p
def num(text):
    p=doc.add_paragraph(style="List Number"); setfont(p.add_run(text)); return p
def para(text,boldlead=None):
    p=doc.add_paragraph()
    if boldlead and text.startswith(boldlead): setfont(p.add_run(boldlead),11,True,DARK); setfont(p.add_run(text[len(boldlead):]))
    else:setfont(p.add_run(text))
    return p

# running furniture
h=sec.header.paragraphs[0]; h.alignment=WD_ALIGN_PARAGRAPH.RIGHT; setfont(h.add_run("地方韌性跨域共創發表會｜世界咖啡館準備文本"),8.5,False,MUTED)
f=sec.footer.paragraphs[0]; f.alignment=WD_ALIGN_PARAGRAPH.CENTER; setfont(f.add_run("2026.07.15｜會前準備與現場發言用"),8.5,False,MUTED)

# cover
p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(56); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; setfont(p.add_run("MEETING BRIEF"),10,True,GOLD)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(7); setfont(p.add_run("地方韌性跨域共創發表會"),27,True,INK)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(22); setfont(p.add_run("世界咖啡館會議準備文本"),16,False,DARK)
table(["日期／時間","地點"],[["7月15日（三）\n13:30–16:30（13:00報到）","兆基文教大樓8樓教室\n臺北市南京東路四段120巷11號"]],[3.25,3.25])
callout("核心任務","不要只盤點補助名稱；協助小組把「地方承受的壓力、政策接不住的地方、韌性架構要求增加的能力」整理成可執行的跨部會建議。")
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; setfont(p.add_run("一個案例｜三個斷點｜三項建議｜兩個問題"),11,True,BLUE)
doc.add_page_break()

doc.add_heading("一、先掌握會議真正要完成的事",1)
callout("一句話判讀","這場會議正在建立地方創生的新政策語言：從單一補助、短期產值與活動人次，轉向地方是否能維持生活、累積關係、調整制度並面對變局。")
para("主辦方安排鹿寮坑溪地、埔里與南區案例，接著以固定五組、不換桌的方式討論兩題，顯示目的不是廣泛交換名片，而是讓跨部會與第一線工作者共同驗證「地方韌性」是否能成為政策診斷、資源組合與成效評估的上位框架。")
doc.add_heading("正式議程",2)
table(["時間","內容","準備焦點"],[("13:35–14:30","案例敘事、成果、規劃報告與工具箱","記下案例如何從產值外說明地方改變"),("14:35–15:45","5組固定討論，不換桌","完成政策資源、挑戰、韌性解法與三項建議"),("15:45–16:20","每組發表7分鐘","用問題—斷點—解法呈現，不逐條念紀錄")],[1.15,2.55,2.8])
doc.add_heading("兩道正式題目",2)
num("各部會目前與地方創生相關的政策資源或補助計畫是什麼？在推動上面臨什麼挑戰？")
num("呼應韌性架構，此架構可以如何解開上述挑戰？")

doc.add_heading("二、三篇閱讀材料形成的共同論述",1)
doc.add_heading("1. 地方韌性是四種能力",2)
table(["能力","判斷問題","政策含義"],[("抵抗力","第一波衝擊時能否守住基本功能？","平時是否有多元收入、備援與基本服務"),("吸收力","受創後能否維持或迅速恢復？","服務中斷多久、誰能支援"),("適應力","條件改變時能否調整運作？","能否改變通路、工作、服務與合作"),("轉型力","能否從危機長出新結構？","是否形成新產業、新治理與新生活方式")],[1.05,2.55,2.9])
callout("評估公式","地方韌性＝結構韌性（地方體質）× 應變韌性（面對衝擊的反應）。因此不能只看現在有多少資源，也要看資源中斷時能否維持、調整與轉型。")

doc.add_heading("2. 六大面向把地方還原為生活系統",2)
table(["面向","最直接的問題"],[("經濟","收入與產業是否多元？錢是否留在地方循環？"),("社會","居民、組織與店家是否互相信任並能動員？"),("人才培育","青年與關鍵工作者能否留下、成長與接班？"),("集體認同","居民是否認為這是自己的家園並願意承擔？"),("自然環境","產業與觀光是否超過環境承載力？"),("基礎建設","交通、醫療、教育、防災與公共空間是否支持生活？")],[1.3,5.2])

doc.add_heading("3. 地方創生是一種政策典範，不只是補助類別",2)
para("延伸文章主張，地方創生若只是青年創業、觀光或產業補助的延續，就沒有獨特政策價值。它應是一個更上位的社會經濟框架，把經濟價值、社會關係、文化、生活與環境放在同一套判斷中。")
bullet("地方不是提供土地、人力、產品與文化素材的「產地」，而是居民生活、工作、照顧與形成認同的場域。")
bullet("不是反對產值，而是追問產值的來源、流向、分配、代價與補助結束後能否延續。")
bullet("地方小微產業與社區網絡像經濟微血管，決定國家財富能否真正流入地方生活。")
bullet("政策應支持能讓收入、人才、關係與能力留在地方的「家園經濟」。")

doc.add_heading("三、第一題可能出現的政策挑戰",1)
table(["可能討論","現場常見表現","深層問題"],[("資源很多但不會組合","同一目標拆成多份計畫","部會依權責分工，地方問題卻是完整系統"),("期程與KPI不一致","各計畫各自結案、重複佐證","沒有共同里程碑與整合窗口"),("重短期成果","偏重場次、人次、曝光、執行率","關係、能力與接班難以被行政語言承認"),("補助結束難以維持","人力退場、空間與設備無人維護","缺少營運收入與長期支持"),("資源集中少數團隊","明星團隊成長，周邊仍脆弱","尚未形成多節點的小微經濟網絡"),("地方價值被外部抽取","外部廠商獲益、成本由居民承擔","產值增加不等於地方受益")],[1.35,2.5,2.65])
callout("可用追問","這項補助創造的收入、人才、合作關係與能力，有多少在結案後仍留在地方？若主力團隊或單一補助退出，地方還能運作嗎？")

doc.add_heading("四、第二題如何用韌性架構解題",1)
num("先說地方承受什麼壓力：人口、高齡、產業、空間、醫療、教育、環境或組織接班。")
num("再指出現有政策在哪裡接不住：資格、期程、KPI、核銷、權責、營運或跨部會斷點。")
num("用六大面向檢查被忽略的系統關係。")
num("用抵抗、吸收、適應、轉型四種能力，說明政策應增加什麼能力。")
num("提出制度改變：整合窗口、資源組合、共同里程碑、共同佐證、長期支持與政策回饋。")
doc.add_heading("政策解法候選",2)
table(["解法","具體作法"],[("從補助導向改為問題導向","先做地方壓力與六面向診斷，再組合部會工具"),("跨部會資源組合","一個地方目標對應多項工具，指定整合窗口"),("共同成果架構","建立共同核心指標、地方自選指標與案例敘事證據"),("分階段長期支持","由能力建立、試行、營運到退場分期，不只單年度補助"),("共同佐證與行政減量","部會共同採認部分資料，避免地方重複填報"),("制度缺口回饋","工具箱診斷出無計畫承接的需求時，能回饋政策調整")],[1.8,4.7])

doc.add_heading("五、最可能出現的三場爭論",1)
doc.add_heading("爭論一：產值還是價值？",2)
callout("建議收斂","不是取消產值，而是追問產值的來源、流向、分配與延續。高度依賴單一補助的高產值未必有韌性；能留在地方、支撐多個角色並持續循環的小規模收入，可能更有韌性。")
doc.add_heading("爭論二：標準化指標還是地方差異？",2)
callout("建議收斂","採「共同核心指標＋地方自選指標＋案例敘事證據」。中央保留比較基礎，地方保留問題脈絡。")
doc.add_heading("爭論三：韌性是地方責任還是政府責任？",2)
callout("建議收斂","韌性不是要求地方自己想辦法，而是重新釐清中央、地方政府、輔導體系與地方組織各自要承擔什麼。交通、醫療、土地、法規與長期財政不能下放給地方團隊。")

doc.add_heading("六、你的角色與可帶入案例",1)
callout("角色定位","地方需求與政府制度之間的轉譯者：把地方完整生活問題，整理成部會可理解的制度斷點、資源組合與驗證方式。")
doc.add_heading("建議案例：公有空間活化",2)
table(["地方需求","涉及系統","可指出的斷點"],[("空間修繕與合法使用","內政、建管、消防、地方建設","硬體完成與營運啟動不同步"),("青年進駐與工作","青年、經濟、勞動、住房","有創業補助但缺生活與長期職涯支持"),("社區照顧與公共服務","衛福、社區、地方政府","服務計畫與空間／交通資源分離"),("文化、遊程與地方內容","文化、觀光、農業、教育","曝光與人次未必形成在地收入循環"),("環境與長期維護","環境、農業、防災、地方政府","建置經費有，維護與退場機制不足")],[1.75,2.15,2.6])
callout("案例結論","地方真正面對的是完整生活系統，但政府資源依部會切割。工具箱應協助地方把完整問題轉譯為可組合的政策工具，而不是再增加一張檢核表。")

doc.add_heading("七、70分鐘討論的實際操作",1)
table(["時間","任務","預期產出"],[("0–10分鐘","成員與政策資源快速盤點","資源名稱、服務對象、政策目的"),("10–30分鐘","整理推動挑戰並合併相似問題","3–5個制度斷點"),("30–50分鐘","用六面向及四能力重新分析","每個斷點缺少的韌性能力"),("50–65分鐘","形成具體政策建議","3項可執行改善"),("65–70分鐘","確認發表主軸與分工","7分鐘報告骨架")],[1.1,2.7,2.7])
doc.add_heading("你可以協助小組使用的四欄",2)
table(["現有資源","推動挑戰","受影響的韌性","政策改善"],[('補助／政策工具','資格、期程、KPI、核銷或權責斷點','六面向＋四能力','整合窗口、共同指標、資源組合或長期機制')],[1.45,1.75,1.45,1.85])

doc.add_heading("八、7分鐘小組發表骨架",1)
table(["時間","內容","說法"],[("1分鐘","政策資源輪廓","本組看到的不是單一計畫，而是若干互補但尚未銜接的工具"),("2分鐘","三個主要挑戰","用地方情境說明制度斷點，不只列行政抱怨"),("2分鐘","韌性架構的重新判讀","說明缺的是哪個面向及哪種能力"),("1.5分鐘","三項政策建議","每項包含誰負責、改什麼、如何驗證"),("0.5分鐘","結論","韌性不是新標籤，而是檢查政策能否讓地方持續運作的框架")],[1.05,1.65,3.8])
callout("小組結論句","地方韌性的價值，不是替既有補助換一個名稱，而是重新檢查：現有政策能否讓地方在資源、人才或環境改變時，仍然維持、調整並形成新的發展能力。")

doc.add_heading("九、可直接使用的發言",1)
doc.add_heading("60秒自我介紹",2)
callout("口語稿","我長期參與政府委辦、地方創生輔導、公有空間活化及地方團隊陪伴。我的觀察是，地方通常不缺創意，也不一定缺補助資訊，真正欠缺的是把完整需求轉譯成跨部會可以共同承接的行動架構。地方韌性不只是災後恢復，而是當補助、人員、市場或環境改變時，地方仍能維持運作、調整方法並形成下一條路。我今天特別想討論的是，工具箱如何協助地方組合資源，而不是再增加一套申請表格。")
doc.add_heading("三個優先問題",2)
num("工具箱診斷出跨部會問題後，誰負責把結果轉成實際資源組合？")
num("共同指標如何保留地方差異，避免韌性變成另一套統一KPI？")
num("如果補助結束後，人力、空間與服務無法維持，是否代表原計畫沒有真正提升韌性？")
doc.add_heading("必要時協助全桌收斂",2)
callout("收斂句","我們可以把討論分成三層：地方承受什麼壓力、現有政策在哪裡接不住、韌性架構要求政策增加什麼能力。這樣就不會只停在補助名稱盤點。")

doc.add_heading("十、會前最後檢查",1)
for x in ["確認行前讀物、工具箱與活動地點。","準備一個公有空間活化或地方團隊案例。","記住三個制度斷點：資源切割、短期KPI、結案後難延續。","選定三項建議：整合窗口、共同成果架構、分階段長期支持。","複習60秒自介及三個優先問題。","現場記錄各部會真正願意調整的事項、分歧與後續聯繫人。"]: bullet("□ "+x)

doc.add_heading("十一、閱讀來源",1)
para("1. 微笑台灣，〈地方韌性是什麼？為什麼它決定了你家鄉的未來？從六大面向打開自我檢視工具箱〉，2026-06-21。https://smiletaiwan.cw.com.tw/article/9210")
para("2. 微笑台灣，〈2026地方韌性數位專輯｜我是全村的希望〉。https://smiletaiwan.cw.com.tw/topics/communityresilience/")
para("3. 謝昇佑，〈產值不是唯一指標，釐清地方創生的政策定位，重點在如何讓地方成為具韌性的共同體〉，2025-10-26。https://smiletaiwan.cw.com.tw/article/8599")
para("4. 「跨域共創發表會」行前通知與正式議程，2026-07-15。")

doc.save(OUT); print(OUT)
