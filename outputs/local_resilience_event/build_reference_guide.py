from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_BREAK

OUT = "/Users/junghanchiu/Documents/New project/outputs/local_resilience_event/地方韌性跨域共創發表會_行前參考指南_20260715.docx"
BLUE = "2E74B5"
DARK = "1F4D78"
INK = "243746"
MUTED = "6B7280"
LIGHT = "E8EEF5"
PALE = "F4F6F9"
GOLD = "A46B12"

doc = Document()
sec = doc.sections[0]
sec.page_width, sec.page_height = Inches(8.5), Inches(11)
sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1)
sec.header_distance = sec.footer_distance = Inches(.492)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "STHeiti"
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "STHeiti")
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.25

for name, size, color, before, after in [
    ("Heading 1", 16, BLUE, 18, 10),
    ("Heading 2", 13, BLUE, 14, 7),
    ("Heading 3", 12, DARK, 10, 5),
]:
    st = styles[name]
    st.font.name = "STHeiti"
    st._element.rPr.rFonts.set(qn("w:eastAsia"), "STHeiti")
    st.font.size = Pt(size)
    st.font.color.rgb = RGBColor.from_string(color)
    st.font.bold = True
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)
    st.paragraph_format.keep_with_next = True

for list_name in ["List Bullet", "List Number"]:
    st = styles[list_name]
    st.font.name = "STHeiti"
    st._element.rPr.rFonts.set(qn("w:eastAsia"), "STHeiti")
    st.font.size = Pt(11)
    st.paragraph_format.left_indent = Inches(.375)
    st.paragraph_format.first_line_indent = Inches(-.188)
    st.paragraph_format.space_after = Pt(4)
    st.paragraph_format.line_spacing = 1.25

if "Guide Callout" not in styles:
    callout = styles.add_style("Guide Callout", WD_STYLE_TYPE.PARAGRAPH)
else:
    callout = styles["Guide Callout"]
callout.font.name = "STHeiti"
callout._element.rPr.rFonts.set(qn("w:eastAsia"), "STHeiti")
callout.font.size = Pt(11)
callout.font.color.rgb = RGBColor.from_string(INK)
callout.paragraph_format.space_before = Pt(6)
callout.paragraph_format.space_after = Pt(8)
callout.paragraph_format.left_indent = Inches(.18)
callout.paragraph_format.right_indent = Inches(.18)
callout.paragraph_format.line_spacing = 1.25

def shade_paragraph(p, fill=PALE):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    pPr.append(shd)
    for side in ["top", "left", "bottom", "right"]:
        borders = pPr.find(qn("w:pBdr"))
        if borders is None:
            borders = OxmlElement("w:pBdr")
            pPr.append(borders)
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), "D9E1EA")
        borders.append(el)

def set_cell(cell, text, bold=False, fill=None, color="000000", align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "STHeiti"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "STHeiti")
    r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    tcPr = cell._tc.get_or_add_tcPr()
    if fill:
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), fill)
        tcPr.append(shd)
    tcMar = OxmlElement("w:tcMar")
    for side, val in [("top", 80), ("bottom", 80), ("start", 120), ("end", 120)]:
        node = OxmlElement(f"w:{side}")
        node.set(qn("w:w"), str(val)); node.set(qn("w:type"), "dxa")
        tcMar.append(node)
    tcPr.append(tcMar)

def add_table(headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for i, h in enumerate(headers):
        set_cell(table.rows[0].cells[i], h, True, LIGHT, DARK, WD_ALIGN_PARAGRAPH.CENTER)
        table.rows[0].cells[i].width = Inches(widths[i])
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell(cells[i], str(value), False, None, "222222")
            cells[i].width = Inches(widths[i])
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    return table

def add_bullet(text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    if level:
        p.paragraph_format.left_indent = Inches(.65)
    p.add_run(text)
    return p

def add_num(text):
    p = doc.add_paragraph(style="List Number")
    p.add_run(text)
    return p

def add_callout(label, text):
    p = doc.add_paragraph(style="Guide Callout")
    r = p.add_run(label + "｜")
    r.bold = True; r.font.color.rgb = RGBColor.from_string(DARK)
    p.add_run(text)
    shade_paragraph(p)

def page_break():
    doc.add_page_break()

# Header/footer
header = sec.header.paragraphs[0]
header.text = "地方韌性跨域共創發表會｜行前參考指南"
header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
header.runs[0].font.size = Pt(8.5); header.runs[0].font.color.rgb = RGBColor.from_string(MUTED)
footer = sec.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer.add_run("2026.07.15｜個人行前準備文件")
run.font.size = Pt(8.5); run.font.color.rgb = RGBColor.from_string(MUTED)

# Cover
doc.add_paragraph().paragraph_format.space_after = Pt(62)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("FIELD GUIDE"); r.bold = True; r.font.size = Pt(10); r.font.color.rgb = RGBColor.from_string(GOLD)
p.paragraph_format.space_after = Pt(14)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("地方韌性跨域共創發表會"); r.bold = True; r.font.size = Pt(27); r.font.color.rgb = RGBColor.from_string(INK)
p.paragraph_format.space_after = Pt(8)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("世界咖啡館參與與政策對話行前參考指南"); r.font.size = Pt(15); r.font.color.rgb = RGBColor.from_string(DARK)
p.paragraph_format.space_after = Pt(26)

add_table(["日期／時間", "地點"], [["2026年7月15日（三）\n13:30–16:30（13:00報到）", "兆基文教大樓8樓教室\n臺北市松山區南京東路四段120巷11號"]], [3.25, 3.25])
doc.add_paragraph().paragraph_format.space_after = Pt(20)
add_callout("核心任務", "理解主辦方的地方韌性框架，帶入政府委辦與地方輔導現場經驗，協助把地方需求轉譯成可被跨部會共同承接的政策路徑。")
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("準備重點：一個案例｜三個斷點｜三項建議｜兩個問題"); r.bold = True; r.font.size = Pt(11); r.font.color.rgb = RGBColor.from_string(BLUE)

page_break()
doc.add_heading("一、快速判讀：這場會議真正要做什麼", level=1)
add_callout("一句話結論", "這不是一般成果發表會，而是研究成果驗證、成效指標蒐集、政策報告轉譯檢討，以及跨部會資源問題盤點的邀請制工作會議。")

doc.add_heading("從報名表看到的四個訊號", level=2)
for t in [
    "主辦方正在尋找比活動場次、參與人次與短期產值更合適的地方創生成效指標。",
    "主辦方關心地方需求如何被轉寫成政策報告，以及執行效益如何被合理敘述。",
    "世界咖啡館可能依資源整合、法規鬆綁、基礎設施、產業與人才等問題分桌。",
    "與會角色集中於中央部會、部會PO、國發會、台經院、分區輔導中心及地方團隊，屬政策執行網絡內部校準。",
]: add_bullet(t)

doc.add_heading("主辦方可能期待的成果", level=2)
for t in [
    "確認地方韌性六大面向與工具箱是否能被實務工作者使用。",
    "建立結果指標、過程指標與韌性驗證的候選指標庫。",
    "辨識現行部會資源、法規、期程與核銷制度的斷點。",
    "形成可寫入研究報告或政策建議的跨部會對接方案。",
]: add_bullet(t)

doc.add_heading("二、地方韌性框架速查", level=1)
add_callout("官方研議定義", "地方面對天災、經濟波動、人口外移、產業衰退等衝擊或長期壓力時，能夠抵抗、吸收、適應並轉型的綜合能力。")
add_table(["能力", "現場理解", "可追問的證據"], [
    ["抵抗力", "第一波衝擊來時能否守住基本功能", "有無備援、人力與資源是否過度集中"],
    ["吸收力", "受創後能否維持或快速恢復", "關鍵服務中斷多久、誰能提供支援"],
    ["適應力", "條件改變時能否調整運作方式", "是否能改變產品、通路、服務或合作方式"],
    ["轉型力", "能否由危機長出新結構", "是否形成新產業、新治理或新生活模式"],
], [1.05, 2.7, 2.75])

doc.add_heading("六大評估面向", level=2)
add_table(["面向", "核心問題"], [
    ["經濟", "地方靠什麼維生？收入與產業是否多元？"],
    ["社會", "居民是否互相認識、信任並能動員？"],
    ["人才培育", "年輕人與關鍵工作者是否留得下來、有人接班？"],
    ["集體認同", "居民是否以地方為榮，願意共同承擔公共事務？"],
    ["自然環境", "地方行動是否超過生態與環境承載力？"],
    ["基礎建設", "交通、醫療、防災及公共設施能否支持生活？"],
], [1.25, 5.25])

page_break()
doc.add_heading("三、你在現場最有價值的角色", level=1)
add_callout("建議定位", "地方需求與政府制度之間的轉譯者：理解地方完整生活系統，也熟悉政府委辦、計畫管理、交付要求與跨部會資源限制。")

doc.add_heading("建議帶入的三個主張", level=2)
claims = [
    ("不要把韌性變成新的補助包裝詞", "韌性應反映收入來源、接班人、替代通路、支援網絡與危機時維持運作的能力，而不是只在計畫名稱加上『韌性』。"),
    ("資源對接不只是補助清單媒合", "應以一個地方韌性目標組合多個部會工具，設定整合窗口、共同里程碑、共同佐證與退場機制。"),
    ("需要地方轉譯者與資源架構師", "分區輔導中心或地方輔導團不應只協助寫計畫，也應協助風險盤點、資源組合、跨部會溝通與成果學習。"),
]
for title, body in claims:
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title); r.bold = True; r.font.color.rgb = RGBColor.from_string(DARK)
    p = doc.add_paragraph(body); p.paragraph_format.left_indent = Inches(.2)

doc.add_heading("可使用的實務案例：公有空間活化", level=2)
add_table(["地方問題", "可能涉及的政策系統"], [
    ["空間修繕與合法使用", "內政、地方建設、建管、消防"],
    ["青年進駐與營運", "青年、經濟、勞動與創業資源"],
    ["社區照顧與公共服務", "衛福、社區發展與地方政府"],
    ["文化、遊程與地方內容", "文化、觀光、農業與教育"],
    ["環境、能源與韌性", "環境、農業、淨零與防災"],
], [2.35, 4.15])
add_callout("案例結論", "地方真正面對的是完整生活系統，但政府資源依部會分割。工具箱的價值，應是把完整問題轉譯成可組合的政策工具，而不是要求地方自己適應政府組織分工。")

doc.add_heading("四、世界咖啡館怎麼參與", level=1)
for t in [
    "每桌約4至8人，圍繞一個核心問題進行短時間討論。",
    "桌長通常留守並整理前一輪內容，其他人換桌延續討論。",
    "最後由各桌提出共同觀察、主要分歧與具體建議。",
]: add_num(t)

doc.add_heading("最有效的發言結構", level=2)
add_callout("發言公式", "一個實際情況 → 一個制度問題 → 一項具體建議。每次控制在一至兩分鐘，讓桌長容易記錄並納入成果。")
add_callout("示範", "我們在公有空間活化時，修繕、進駐、人才與營運分屬不同政策。地方必須把同一目標拆成數份計畫，期程與KPI也無法銜接。建議工具箱增加跨部會資源組合表，並明定整合窗口。")

page_break()
doc.add_heading("五、報名表建議填答稿", level=1)
doc.add_heading("1. 最合適的地方創生政策效益指標", level=2)
doc.add_paragraph("建議依地方原始條件及主要風險設定基準線，衡量地方行動前後的變化，包括地方收入來源多元程度、在地就業與人才留任時間、青年及核心工作者接班情形、地方組織間持續合作關係、居民參與地方公共事務的頻率、關鍵生活服務可及性、地方據點持續營運能力，以及遭遇市場、人口或災害衝擊後維持運作與恢復的時間。除結果指標外，也應納入關係建立、組織學習及跨域協作等過程指標。")

doc.add_heading("2. 最不合適的地方創生政策效益指標", level=2)
doc.add_paragraph("不宜單獨使用活動場次、一次性參與人次、媒體曝光量、補助經費執行率、短期營業額、單年度新增工作數及短期人口增減作為主要判斷。這些數字容易受到地方規模、區位與資源差異影響，也可能鼓勵團隊追求容易量化的活動，而忽略組織能力、社會關係與長期營運。")

doc.add_heading("3. 建議勾選項目", level=2)
add_table(["表單題目", "建議選擇"], [
    ["最怕的報告環節", "整體政策／地方團隊執行效益；地方議題與需求盤點"],
    ["地方創生的政策屬性", "人口、國土、產業發展、社區營造；視需要加文化政策"],
    ["最需跨部會協調", "資源與補助整合、法規鬆綁、產業發展與人才留育；公有空間案例可加基礎設施"],
    ["單位屬性", "若非代表分區輔導中心出席，以『地方創生團隊』最接近"],
], [2.05, 4.45])

doc.add_heading("4. 最大挑戰或痛點", level=2)
doc.add_paragraph("地方實際面對的是人口、產業、空間、社福、文化與環境相互連動的完整問題，但中央資源依部會權責、年度預算與補助科目切割。地方團隊必須將同一發展目標拆成多份計畫，承受不同期程、資格、KPI、核銷及成果格式，卻缺少跨部會整合窗口與長期營運經費。建議由輔導體系協助建立共同問題診斷、資源組合路徑、共同里程碑與成果指標。")

doc.add_heading("5. 希望深度交流的部會", level=2)
add_callout("首選", "內政部國土、城鄉發展及公有空間活化相關業務。希望討論地方創生據點在土地使用、建管消防、空間修繕、營運管理與長期維護之間，如何建立跨計畫銜接機制。")
add_callout("第二選擇", "勞動部多元就業與培力就業業務，討論補助型人力如何轉化為地方長期人才與組織能力。")

page_break()
doc.add_heading("六、可直接使用的現場發言", level=1)
doc.add_heading("60秒自我介紹", level=2)
add_callout("口語稿", "我長期參與政府委辦計畫、地方創生輔導、公有空間活化及地方團隊陪伴。我的觀察是，地方通常不缺創意，也不一定缺補助資訊，真正欠缺的是把地方完整需求轉譯成跨部會可以共同承接的行動架構。對我而言，地方韌性不只是災後恢復，而是當補助、人員、市場或環境條件改變時，地方仍有能力維持運作、調整方法並形成下一條路。我今天特別想討論的是，這套工具箱如何協助地方組合資源，而不是再增加一套申請表格。")

doc.add_heading("一句關鍵主張", level=2)
add_callout("建議發言", "地方韌性工具箱不能只協助地方證明自己有問題，也必須協助中央看見：這個問題需要哪些部會共同承接、誰負責整合，以及補助結束後由什麼機制維持。")

doc.add_heading("六個可主動提出的問題", level=2)
questions = [
    "工具箱主要使用者是地方團隊、地方政府、輔導中心，還是中央部會？不同使用者需要的工具是否應分版？",
    "工具箱會不會提供跨部會資源的組合邏輯，而不只是資源清單？",
    "除了產值與人次，是否納入組織接班、收入多元性、合作網絡、基本服務與恢復時間？",
    "如果盤點出真正需求，但目前沒有補助科目可以承接，是否有回饋政策調整的機制？",
    "跨部會計畫能否建立共同成果指標與佐證格式，降低地方重複行政？",
    "韌性需要長期累積，但補助通常一至三年，後續營運與維護成本由誰承接？",
]
for q in questions: add_num(q)

doc.add_heading("建議工具：政策資源鑲嵌表", level=2)
add_table(["地方壓力", "韌性能力", "地方行動", "政策對接"], [
    ["人口外移／缺工", "適應、轉型", "人才培育、接班與多元工作", "勞動、青年、教育、經濟"],
    ["據點難以維持", "抵抗、吸收", "空間合法化與營運收入", "內政、地方政府、經濟"],
    ["產業單一", "抵抗、轉型", "多元通路與跨域合作", "經濟、農業、觀光、文化"],
], [1.55, 1.25, 1.8, 1.9])

doc.add_heading("七、行前與當日檢查表", level=1)
for t in [
    "完成線上報名，確認單位屬性與希望交流的部會。",
    "讀完《地方韌性工具箱》與微笑台灣地方韌性專欄。",
    "準備一個成功案例及一個跨部會卡關案例。",
    "選定三項制度建議與最想問的兩個問題。",
    "複習60秒自我介紹，避免現場長篇說明背景。",
    "攜帶名片、手機／平板、充電設備及本參考指南。",
    "13:00至13:15抵達兆基文教大樓8樓；注意不是國發會辦公大樓。",
    "記下重要與會者、可合作對象、政策關鍵詞與後續聯繫事項。",
]: add_bullet("☐ " + t)

doc.add_heading("八、建議會前向謝昇佑確認", level=1)
add_callout("可直接傳訊", "這次最希望我從『地方執行現場、政府委辦管理、成效指標或跨部會整合』哪一個角度提供回饋？如果世界咖啡館已有預定桌次，也請告訴我，我可以事先準備更具體的案例。")

doc.add_heading("九、資料來源", level=1)
sources = [
    "活動報名表：114年度《地方創生實踐行動價值研析與深化》計畫｜共創發表會｜報名表。",
    "國發會地方創生政策：打造永續共好地方創生計畫（114–117年）。https://www.ndc.gov.tw/nc_16183_40004",
    "國發會地方創生政策緣起。https://www.twrr.ndc.gov.tw/policyOrigin",
    "《地方韌性是什麼？為什麼它決定了你家鄉的未來？從六大面向打開自我檢視工具箱》。",
    "《地方韌性工具箱》，國發會《地方創生實踐行動價值研析與深化》地方創生×地方韌性研議小組。",
]
for s in sources: add_bullet(s)

doc.core_properties.title = "地方韌性跨域共創發表會行前參考指南"
doc.core_properties.subject = "2026年7月15日世界咖啡館與政策對話準備"
doc.core_properties.author = "Codex"
doc.save(OUT)
print(OUT)
