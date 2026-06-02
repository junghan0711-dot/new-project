from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = "/Users/junghanchiu/Documents/New project/outputs/新豐鄉公所公共空間合約書優化/新豐鄉公所公共空間合約書_精簡兩頁版.docx"


def set_run(run, size=11, bold=False, color=None, font="標楷體"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def pformat(p, before=0, after=3, line=1.05, align=None):
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    if align is not None:
        p.alignment = align


def para(doc, text="", size=11, bold=False, before=0, after=3, align=None, color=None):
    p = doc.add_paragraph()
    pformat(p, before=before, after=after, align=align)
    r = p.add_run(text)
    set_run(r, size=size, bold=bold, color=color)
    return p


def clause(doc, title, body):
    p = para(doc, title, size=11.5, bold=True, before=4, after=1, color="1F4D78")
    b = para(doc, body, size=10.8, after=2)
    b.paragraph_format.first_line_indent = Inches(0.22)
    return p


def cell_text(cell, text, bold=False, fill=None, size=10.5):
    cell.text = ""
    p = cell.paragraphs[0]
    pformat(p, after=0, line=1.0)
    r = p.add_run(text)
    set_run(r, size=size, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if fill:
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), fill)
        tc_pr.append(shd)


def set_table(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width
            tc_pr = row.cells[idx]._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width.inches * 1440)))
            tc_w.set(qn("w:type"), "dxa")


doc = Document()
sec = doc.sections[0]
sec.page_width = Inches(8.27)
sec.page_height = Inches(11.69)
sec.top_margin = Inches(0.65)
sec.bottom_margin = Inches(0.65)
sec.left_margin = Inches(0.75)
sec.right_margin = Inches(0.75)

normal = doc.styles["Normal"]
normal.font.name = "標楷體"
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "標楷體")
normal.font.size = Pt(11)

title = doc.add_paragraph()
pformat(title, after=2, line=1.05, align=WD_ALIGN_PARAGRAPH.CENTER)
r = title.add_run("新竹縣新豐鄉公所新庄子公有零售市場二樓創生基地\n公共空間借用（使用）契約書")
set_run(r, size=15, bold=True, color="0B2545")
para(doc, "精簡範本｜適用於進駐商家、合作單位或專案推廣對象", size=9.5, after=5, align=WD_ALIGN_PARAGRAPH.CENTER, color="555555")

intro = (
    "新竹縣新豐鄉公所（以下簡稱甲方）與＿＿＿＿＿＿＿＿＿＿＿＿＿＿（以下簡稱乙方），"
    "就乙方借用新庄子公有零售市場二樓創生基地公共空間事宜，經雙方同意訂立本契約。"
)
para(doc, intro, size=10.8, after=4)

para(doc, "一、契約基本資料", size=11.5, bold=True, color="1F4D78", after=2)
tbl = doc.add_table(rows=7, cols=2)
tbl.style = "Table Grid"
set_table(tbl, [Inches(1.15), Inches(5.87)])
rows = [
    ("甲乙雙方", "甲方：新竹縣新豐鄉公所　　乙方：＿＿＿＿＿＿＿＿＿＿＿＿＿＿"),
    ("借用場地", "新庄子公有零售市場二樓創生基地公共空間"),
    ("使用期間", "中華民國＿＿年＿＿月＿＿日起至＿＿年＿＿月＿＿日止；每日＿＿時至＿＿時"),
    ("使用目的", "＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿"),
    ("保證金", "新臺幣＿＿＿＿＿＿元整"),
    ("使用費", "新臺幣＿＿＿＿＿＿元整；每日新臺幣＿＿＿＿＿＿元，依甲方公告或核定為準"),
    ("優惠資格", "□ 適用免收/優惠　　□ 不適用　　聯絡窗口：＿＿＿＿＿＿＿＿＿＿"),
]
for row, data in zip(tbl.rows, rows):
    cell_text(row.cells[0], data[0], bold=True, fill="F2F4F7")
    cell_text(row.cells[1], data[1], size=10.5)

clause(doc, "二、使用範圍與期間", "乙方應依甲方核准之使用目的、期間、時段及場地範圍使用本場地；未經甲方書面同意，不得變更用途、活動內容、展售品項、場地配置或延長使用期間。")
clause(doc, "三、費用、保證金與優惠", "乙方應於甲方指定期限內繳納保證金及使用費。經甲方認定為進駐商家、合作單位或專案推廣對象者，自首次核准使用日起一年內得免收場地使用費；第二年起依甲方公告收費標準計收。乙方違反本契約或管理規定者，甲方得取消優惠資格並追繳原應繳納之使用費。")
clause(doc, "四、乙方應遵守事項", "乙方不得將場地轉借、分租或變相提供他人使用；不得擅自變更活動內容；不得使車輛進入本場地。乙方應遵守零售市場管理條例、新竹縣公有零售市場自治組織設置辦法、甲方公告及相關法令，並維持公共安全、環境清潔及市場秩序。使用結束後，乙方應立即清除器物、垃圾與自備設備，並回復原狀。")
clause(doc, "五、損害賠償與責任", "乙方應以善良管理人之注意義務使用本場地及設備。借用期間如造成場地、設施、植栽、公共財物或第三人損害，或因活動發生消費、商品、食品、智慧財產權、公共安全等糾紛，均由乙方自行負責；甲方如受損害或遭第三人請求，乙方應負賠償及處理責任。")
clause(doc, "六、甲方管理權與停止使用", "甲方基於公共安全、市場管理、公務需求、重大活動或公共利益，得要求乙方改善、調整使用、暫停或停止使用。乙方違約且未改善，或情節重大、急迫危及公共安全或市場秩序者，甲方得立即停止乙方使用並沒入保證金；乙方不得請求補償。")

doc.add_page_break()

clause(doc, "七、取消使用與退費", "因天災、事變、疫情、法令變更、重大公共安全事故、公務使用、重大活動或其他不可歸責於雙方之事由，致乙方無法使用本場地者，甲方得撤銷或變更核准使用，並無息退還未使用部分之保證金及使用費；已使用日數之使用費不予退還。因可歸責於乙方之事由而取消、停止或未使用者，已繳使用費不予退還。")
clause(doc, "八、保證金退還", "乙方使用期滿，經甲方確認無違約、無待清潔、待修復、待賠償或其他應付款項後，甲方依規定無息退還保證金；如有應扣款項，甲方得自保證金中扣抵，不足部分乙方仍應補足。")
clause(doc, "九、未盡事宜與管轄法院", "本契約未盡事項，依零售市場管理條例、新竹縣公有零售市場自治組織設置辦法、甲方公告及相關法令辦理；必要時雙方得以書面補充協議。因本契約所生爭議，雙方同意以臺灣新竹地方法院為第一審管轄法院。")
clause(doc, "十、契約份數", "本契約一式二份，甲乙雙方各執一份為憑；附件、核准文件、點交紀錄、甲方公告或補充協議，均視為本契約之一部分。")

para(doc, "立契約書人", size=12, bold=True, before=6, after=4, align=WD_ALIGN_PARAGRAPH.CENTER)
sign = doc.add_table(rows=7, cols=2)
sign.style = "Table Grid"
set_table(sign, [Inches(3.51), Inches(3.51)])
sign_rows = [
    ("甲　　方：新竹縣新豐鄉公所", "乙　　方：＿＿＿＿＿＿＿＿＿＿"),
    ("代表人：＿＿＿＿＿＿＿＿＿＿", "代表人：＿＿＿＿＿＿＿＿＿＿"),
    ("統一編號：＿＿＿＿＿＿＿＿", "統一編號：＿＿＿＿＿＿＿＿"),
    ("地　　址：＿＿＿＿＿＿＿＿＿＿＿＿＿＿", "地　　址：＿＿＿＿＿＿＿＿＿＿＿＿＿＿"),
    ("聯絡電話：＿＿＿＿＿＿＿＿", "聯絡電話：＿＿＿＿＿＿＿＿"),
    ("電子郵件：＿＿＿＿＿＿＿＿", "電子郵件：＿＿＿＿＿＿＿＿"),
    ("簽章：", "簽章："),
]
for row, vals in zip(sign.rows, sign_rows):
    for idx, text in enumerate(vals):
        cell_text(row.cells[idx], text, size=10.6)

para(doc, "中　華　民　國＿＿年＿＿月＿＿日", size=11, before=8, after=0, align=WD_ALIGN_PARAGRAPH.CENTER)

footer = sec.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = footer.add_run("新竹縣新豐鄉公所公共空間借用（使用）契約書精簡範本")
set_run(fr, size=8.5, color="666666")

doc.save(OUTPUT)
print(OUTPUT)
