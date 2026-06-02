from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = "/Users/junghanchiu/Documents/New project/outputs/新豐鄉公所公共空間合約書優化/新豐鄉公所公共空間合約書_優化版.docx"


def set_run_font(run, size=12, bold=False, color=None, name="標楷體"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_para_format(paragraph, before=0, after=6, line=1.2, align=None):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line
    if align is not None:
        paragraph.alignment = align


def add_paragraph(doc, text="", size=12, bold=False, after=6, before=0, align=None, color=None):
    p = doc.add_paragraph()
    set_para_format(p, before=before, after=after, align=align)
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold, color=color)
    return p


def add_clause_heading(doc, title):
    p = add_paragraph(doc, title, size=13, bold=True, before=8, after=4, color="1F4D78")
    return p


def add_clause_text(doc, text):
    p = add_paragraph(doc, text, size=12, after=4)
    p.paragraph_format.first_line_indent = Inches(0.28)
    return p


def add_item(doc, label, text):
    p = doc.add_paragraph()
    set_para_format(p, after=3, line=1.18)
    p.paragraph_format.left_indent = Inches(0.24)
    p.paragraph_format.first_line_indent = Inches(-0.24)
    r1 = p.add_run(label)
    set_run_font(r1, bold=True)
    r2 = p.add_run(text)
    set_run_font(r2)
    return p


def set_cell_text(cell, text, bold=False, fill=None, width=None):
    cell.text = ""
    p = cell.paragraphs[0]
    set_para_format(p, after=0, line=1.15)
    r = p.add_run(text)
    set_run_font(r, size=11, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if fill:
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), fill)
        tc_pr.append(shd)
    if width is not None:
        cell.width = width


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            if idx < len(row.cells):
                row.cells[idx].width = width
                tc_pr = row.cells[idx]._tc.get_or_add_tcPr()
                tc_w = tc_pr.first_child_found_in("w:tcW")
                if tc_w is None:
                    tc_w = OxmlElement("w:tcW")
                    tc_pr.append(tc_w)
                tc_w.set(qn("w:w"), str(int(width.inches * 1440)))
                tc_w.set(qn("w:type"), "dxa")


def shade_header(row, fill="E8EEF5"):
    for cell in row.cells:
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), fill)
        tc_pr.append(shd)


def style_table_text(table):
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                set_para_format(p, after=0, line=1.15)
                for run in p.runs:
                    set_run_font(run, size=11)


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(1.0)
section.bottom_margin = Inches(1.0)
section.left_margin = Inches(1.25)
section.right_margin = Inches(1.25)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "標楷體"
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "標楷體")
normal.font.size = Pt(12)

title = doc.add_paragraph()
set_para_format(title, after=4, line=1.15, align=WD_ALIGN_PARAGRAPH.CENTER)
run = title.add_run("新竹縣新豐鄉公所新庄子公有零售市場二樓創生基地\n公共空間借用（使用）契約書")
set_run_font(run, size=17, bold=True, color="0B2545")

subtitle = add_paragraph(doc, "範本｜適用於進駐商家、合作單位或專案推廣對象借用公共空間", size=10.5, after=10, align=WD_ALIGN_PARAGRAPH.CENTER, color="555555")

intro = (
    "新竹縣新豐鄉公所（以下簡稱甲方）與＿＿＿＿＿＿＿＿＿＿＿＿＿＿（以下簡稱乙方），"
    "就乙方借用新庄子公有零售市場二樓創生基地公共空間事宜，經雙方同意訂立本契約，"
    "以明確使用範圍、費用、管理規範及雙方權利義務。"
)
add_paragraph(doc, intro, after=8)

add_paragraph(doc, "契約基本資料", size=13, bold=True, before=4, after=4, color="1F4D78")
basic = doc.add_table(rows=12, cols=2)
basic.style = "Table Grid"
set_table_geometry(basic, [Inches(1.25), Inches(4.75)])
labels = [
    ("甲方", "新竹縣新豐鄉公所"),
    ("乙方", "＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿"),
    ("借用場地", "新庄子公有零售市場二樓創生基地公共空間"),
    ("使用目的", "＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿"),
    ("使用期間", "中華民國＿＿年＿＿月＿＿日起至＿＿年＿＿月＿＿日止"),
    ("使用時段", "每日＿＿時至＿＿時"),
    ("保證金", "新臺幣＿＿＿＿＿＿元整"),
    ("使用費", "新臺幣＿＿＿＿＿＿元整"),
    ("計費方式", "每日新臺幣＿＿＿＿＿＿元；實際金額依甲方公告或核定為準"),
    ("優惠資格", "□ 適用　　□ 不適用"),
    ("聯絡窗口", "甲方：＿＿＿＿＿＿＿＿＿＿＿＿＿＿乙方：＿＿＿＿＿＿＿＿＿＿＿＿＿＿"),
    ("備註", "＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿"),
]
for row, data in zip(basic.rows, labels):
    for idx, text in enumerate(data):
        is_label = idx == 0
        set_cell_text(row.cells[idx], text, bold=is_label, fill="F2F4F7" if is_label else None)
style_table_text(basic)

add_clause_heading(doc, "第一條　契約目的與場地")
add_clause_text(doc, "甲方同意乙方依本契約及甲方核准之用途，借用新庄子公有零售市場二樓創生基地公共空間（以下簡稱本場地）。本場地之實際位置、可使用範圍、設備及限制事項，依甲方核准文件、現場公告或點交紀錄為準。")

add_clause_heading(doc, "第二條　使用期間與用途")
add_item(doc, "一、", "乙方使用期間為中華民國＿＿年＿＿月＿＿日起至＿＿年＿＿月＿＿日止；每日使用時段以甲方核准內容為準。")
add_item(doc, "二、", "乙方應依核准之使用目的、活動內容及場地配置使用本場地。未經甲方書面同意，不得擅自變更用途、活動內容、展售品項、使用範圍或現場布置。")
add_item(doc, "三、", "乙方如需延長使用期間或變更使用內容，應於使用日前向甲方提出申請，經甲方同意後始得辦理。")

add_clause_heading(doc, "第三條　費用、保證金與繳納")
add_item(doc, "一、", "保證金為新臺幣＿＿＿＿＿＿元整；使用費為新臺幣＿＿＿＿＿＿元整。實際收費標準依甲方公告、核准文件或雙方約定辦理。")
add_item(doc, "二、", "乙方應於甲方指定期限內繳納保證金及使用費；逾期未繳納者，甲方得取消核准使用，乙方不得請求任何補償。")
add_item(doc, "三、", "保證金不得主張抵充使用費、損害賠償或其他應付款項，但甲方依本契約抵扣或沒入者，不在此限。")

add_clause_heading(doc, "第四條　免收或優惠使用費")
add_item(doc, "一、", "乙方經甲方認定為進駐商家、合作單位或專案推廣對象者，自首次核准使用日起一年內，得免收場地使用費。")
add_item(doc, "二、", "前款優惠期間屆滿後，乙方如續行使用本場地，第二年起依甲方公告收費標準或核准條件計收使用費。")
add_item(doc, "三、", "乙方於優惠期間內違反本契約、甲方公告事項或相關管理規定者，甲方得取消優惠資格，並得追繳原應繳納之使用費。")

add_clause_heading(doc, "第五條　乙方應遵守事項")
rules = [
    "不得將本場地轉借、分租、提供他人使用，或以其他方式變相移轉使用權。",
    "不得擅自變更活動內容、營業或展示品項、場地配置、用電方式、廣告招牌或其他經甲方核准之事項。",
    "不得違反零售市場管理條例、新竹縣公有零售市場自治組織設置辦法、甲方公告之管理規範及其他相關法令。",
    "不得使車輛進入本場地；如因搬運、施工或特殊需要，應事先取得甲方同意並依指定動線、時間及方式辦理。",
    "應維持公共安全、消防安全、用電安全、環境清潔、噪音管制及公共秩序，不得妨害市場營運、鄰近攤商或其他使用人之權益。",
    "活動或使用結束後，應立即清除場地內之器物、垃圾、布置物及其他自備設備，並將場地回復原狀。"
]
for idx, rule in enumerate(rules, start=1):
    add_item(doc, f"{idx}.", rule)

add_clause_heading(doc, "第六條　場地維護與損害賠償")
add_item(doc, "一、", "乙方應以善良管理人之注意義務使用本場地及相關設施、設備、植栽與公共財物。")
add_item(doc, "二、", "使用期間如因乙方、乙方人員、受邀人、承包廠商或參與活動人員之行為，致本場地、設施、設備、植栽或第三人權益受損，乙方應負修復、賠償及相關法律責任。")
add_item(doc, "三、", "甲方得就修復費用、清潔費、賠償費或其他應付款項，自保證金中扣抵；保證金不足者，乙方仍應補足。")

add_clause_heading(doc, "第七條　活動安全、糾紛與法律責任")
add_clause_text(doc, "乙方因借用本場地舉辦活動、展示、展售、推廣或其他使用行為所生之民事、刑事、行政責任、消費糾紛、食品或商品責任、智慧財產權爭議、公共安全事件及第三人損害，均由乙方自行負責處理；如致甲方受有損害或遭第三人請求，乙方應負賠償及處理責任。")

add_clause_heading(doc, "第八條　甲方管理權與停止使用")
add_item(doc, "一、", "甲方基於公共安全、市場管理、公務需求、重大活動、法令要求或維護公共利益之必要，得指示乙方改善、限制使用、暫停使用或調整使用範圍。")
add_item(doc, "二、", "乙方違反本契約或相關管理規定，經甲方通知限期改善而未改善，或情節重大、急迫危及公共安全或市場秩序者，甲方得立即停止乙方使用並沒入保證金；乙方因此所受損失，應自行負責，不得異議。")

add_clause_heading(doc, "第九條　取消使用、不可抗力與費用退還")
add_item(doc, "一、", "本契約訂立後，如因天災、事變、疫情、法令變更、重大公共安全事故、甲方重大活動、公務使用或其他不可歸責於雙方之事由，致乙方無法使用本場地者，甲方得撤銷或變更核准使用，並無息退還未使用部分之保證金及使用費。")
add_item(doc, "二、", "乙方已實際使用之日數或時段，其使用費不予退還。但甲方另有公告或雙方另有書面約定者，從其約定。")
add_item(doc, "三、", "如因可歸責於乙方之事由取消、停止或未使用本場地，已繳使用費不予退還；甲方如另受損害，乙方仍應負賠償責任。")

add_clause_heading(doc, "第十條　保證金退還")
add_clause_text(doc, "乙方使用期滿且未違反本契約、未有待清潔、待修復、待賠償或其他應付款項者，甲方於完成場地確認及相關程序後，依規定無息退還保證金。")

add_clause_heading(doc, "第十一條　未盡事宜")
add_clause_text(doc, "本契約未盡事項，依零售市場管理條例、新竹縣公有零售市場自治組織設置辦法、甲方公告之場地管理規範及其他相關法令辦理；必要時，甲乙雙方得以書面補充協議之。")

add_clause_heading(doc, "第十二條　管轄法院")
add_clause_text(doc, "因本契約所生之一切爭議，雙方應本誠信原則協商解決；協商不成而涉訟時，雙方同意以臺灣新竹地方法院為第一審管轄法院。")

add_clause_heading(doc, "第十三條　契約份數")
add_clause_text(doc, "本契約書一式二份，甲乙雙方各執一份為憑；如有附件或補充協議，均視為本契約之一部分。")

doc.add_page_break()
add_paragraph(doc, "立契約書人", size=14, bold=True, after=8, align=WD_ALIGN_PARAGRAPH.CENTER)
sign = doc.add_table(rows=7, cols=2)
sign.style = "Table Grid"
set_table_geometry(sign, [Inches(3.0), Inches(3.0)])
sign_data = [
    ("甲　　方：新竹縣新豐鄉公所", "乙　　方：＿＿＿＿＿＿　　　"),
    ("代表人：＿＿＿＿＿＿　　　", "代表人：＿＿＿＿＿＿　　　"),
    ("統一編號：＿＿＿＿＿＿　　", "統一編號：＿＿＿＿＿＿　　"),
    ("地　　址：＿＿＿＿＿＿　　", "地　　址：＿＿＿＿＿＿　　"),
    ("聯絡電話：＿＿＿＿＿＿　　", "聯絡電話：＿＿＿＿＿＿　　"),
    ("電子郵件：＿＿＿＿＿＿　　", "電子郵件：＿＿＿＿＿＿　　"),
    ("用印：", "用印："),
]
for row, data in zip(sign.rows, sign_data):
    for idx, text in enumerate(data):
        set_cell_text(row.cells[idx], text)
        row.cells[idx].height = Inches(0.38)
style_table_text(sign)

add_paragraph(doc, "", after=8)
date_p = add_paragraph(doc, "中　華　民　國＿＿年＿＿月＿＿日", size=12, after=14, align=WD_ALIGN_PARAGRAPH.CENTER)

add_paragraph(doc, "附件一：場地點交及退場檢核表（可選用）", size=13, bold=True, before=8, after=6, color="1F4D78")
check = doc.add_table(rows=1, cols=4)
check.style = "Table Grid"
set_table_geometry(check, [Inches(0.5), Inches(2.55), Inches(1.25), Inches(1.7)])
hdr = check.rows[0].cells
for cell, text in zip(hdr, ["項次", "檢核項目", "結果", "備註"]):
    set_cell_text(cell, text, bold=True, fill="E8EEF5")
items = [
    "場地及公共走道已回復原狀",
    "垃圾、布置物及自備設備已清除",
    "桌椅、燈具、插座及既有設備無損壞",
    "牆面、地面、門窗及公共設施無污損",
    "消防、用電及安全設備未遭遮蔽或破壞",
    "植栽、景觀及公共財物無損害",
    "鑰匙、門禁卡或其他交付物已返還",
    "其他需改善事項已完成",
]
for idx, item in enumerate(items, start=1):
    cells = check.add_row().cells
    set_cell_text(cells[0], str(idx))
    set_cell_text(cells[1], item)
    set_cell_text(cells[2], "□ 合格　□ 待改善")
    set_cell_text(cells[3], "")
style_table_text(check)

add_paragraph(doc, "甲方點交人：＿＿＿＿＿＿　　　　　乙方確認人：＿＿＿＿＿＿　　　　　日期：＿＿年＿＿月＿＿日", size=11, before=6, after=4)

footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = footer.add_run("新竹縣新豐鄉公所公共空間借用（使用）契約書範本")
set_run_font(fr, size=9, color="666666")

doc.save(OUTPUT)
print(OUTPUT)
